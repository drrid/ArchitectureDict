from requests import get
from bs4 import BeautifulSoup
from docx import Document

doc = Document()
DICT_BASE_URL = "http://www.linternaute.com"
all_terms = []

def get_terms(pageNumber):
	url = "{}/dictionnaire/fr/theme/architecture/{}/".format(DICT_BASE_URL, pageNumber)
	base_url = get(url)
	soup = BeautifulSoup(base_url.content, "html.parser")

	term_grid = soup.find("ul", class_="dico_liste grid_line")
	terms = term_grid.find_all("li")
	for term in terms:
		all_terms.append(DICT_BASE_URL + term.a.get("href"))


def get_term(term):
	n = 0
	base_url = get(term)
	soup = BeautifulSoup(base_url.content, "html.parser")
	for l, el  in enumerate(soup.find_all("div", class_="grid_left")):
		if el.div and ('Architecture' in el.div.get_text()):
			n = l

	definition = soup.find_all("div", class_="grid_last")[n].get_text().strip()
	definition = ' '.join(definition.split())

	title = ' '.join(soup.find("span", class_="dico_title_2").get_text().split())

	return [title, definition]


if __name__ == '__main__':

	for i in ["1", "2", "3"]:
		get_terms(i)

	for index, t in enumerate(all_terms):
		term_list = get_term(t)

		title = term_list[0].split(',')[0]
		catg = term_list[0].split(',')[1].strip()
		content = term_list[1]


		doc.add_heading(title + ' :', level=1)
		doc.add_heading(catg, level=3)
		doc.add_paragraph(term_list[1])
		print str(index) + ' . ' + title

	doc.save('dict_test.docx')

